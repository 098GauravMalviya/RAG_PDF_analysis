import os
import json
import shutil
import time 
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, Any

# Import analyzer components
from analyzer import TenderAnalyzer, BiddingTemplate, CompanyProfile

app = FastAPI(
    title="L&T Defence Tender Analyzer Backend",
    description="FastAPI service for Stage 1-3 Hybrid Tender Analysis and Word report generation.",
    version="1.0"
)

# Enable CORS for frontend communications
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROFILE_PATH = os.path.join(BASE_DIR, "company_profile.json")
OUTPUT_PATH = os.path.join(BASE_DIR, "bidding_template_output.json")
TEMP_DIR = os.path.join(BASE_DIR, "temp_uploads")

# Ensure temp directory exists
os.makedirs(TEMP_DIR, exist_ok=True)

# Shared memory/in-memory task tracker
analysis_tasks: Dict[str, Dict[str, Any]] = {}

# ==========================================
# 1. Helper: Docx Generation Logic
# ==========================================
def build_docx_report(result: BiddingTemplate, filepath: str):
    """Compiles the hybrid scorecard into a professional styled MS Word document."""
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        raise ImportError("python-docx is not installed or available.")

    doc = docx.Document()

    # Define color scheme (navy primary, dark grey body, red/orange/green alerts)
    COLOR_PRIMARY = RGBColor(12, 35, 64)   # Navy
    COLOR_SECONDARY = RGBColor(112, 128, 144) # Slate grey
    COLOR_DARK = RGBColor(51, 51, 51)
    
    # helper for styling tables
    def set_cell_background(cell, color_hex):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
        
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Document Title
    title = doc.add_paragraph()
    title_run = title.add_run("L&T Defence Bidding & Compliance Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Tender Reference: {result.tender_id} | Issued by: {result.issuing_authority}")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = COLOR_SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Recommendation Callout Box
    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell = rec_table.rows[0].cells[0]
    set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
    
    rec_color = "D4EDDA" # green background
    border_color = "28A745"
    rec_text = f"OVERALL BID DECISION: {result.recommendation.value}"
    
    if result.recommendation.value == "NO_GO":
        rec_color = "F8D7DA" # red
        border_color = "DC3545"
    elif result.recommendation.value == "CONDITIONAL_GO":
        rec_color = "FFF3CD" # yellow
        border_color = "FFC107"
        
    set_cell_background(cell, rec_color)
    
    p = cell.paragraphs[0]
    prun = p.add_run(rec_text)
    prun.font.bold = True
    prun.font.size = Pt(14)
    prun.font.color.rgb = COLOR_PRIMARY
    
    p_rat = cell.add_paragraph()
    p_rat_run = p_rat.add_run(f"Rationale: {result.recommendation_rationale}")
    p_rat_run.font.size = Pt(10)
    p_rat_run.font.color.rgb = COLOR_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 1: Scope of Work
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Scope of Work")
    h1_run.font.color.rgb = COLOR_PRIMARY
    h1_run.font.size = Pt(14)
    
    scope_p = doc.add_paragraph(result.scope_of_work)
    scope_p.style.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section 2: Deterministic Scorecard (Stage 2 Python Checks)
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Financial & Eligibility Check (Deterministic)")
    h2_run.font.color.rgb = COLOR_PRIMARY
    h2_run.font.size = Pt(14)

    det_table = doc.add_table(rows=1, cols=4)
    det_table.style = 'Table Grid'
    hdr_cells = det_table.rows[0].cells
    headers = ["Parameter / Check", "Required spec", "L&T Value", "Status"]
    for i, name in enumerate(headers):
        hdr_cells[i].text = name
        set_cell_background(hdr_cells[i], "0C2340") # Navy header
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)

    for check in result.deterministic_scorecard:
        row_cells = det_table.add_row().cells
        row_cells[0].text = check.parameter
        row_cells[1].text = check.required_value
        row_cells[2].text = check.company_value
        
        status_cell = row_cells[3]
        status_cell.text = check.status.value
        set_cell_margins(status_cell, top=80, bottom=80, left=100, right=100)
        
        # Color match status cell
        if check.status.value == "COMPLIANT":
            set_cell_background(status_cell, "D4EDDA")
        else:
            set_cell_background(status_cell, "F8D7DA")
            
        for c in row_cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].style.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 3: Technical specifications scorecard (Stage 3 LLM checks)
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Technical & Engineering Scorecard")
    h3_run.font.color.rgb = COLOR_PRIMARY
    h3_run.font.size = Pt(14)

    tech_table = doc.add_table(rows=1, cols=4)
    tech_table.style = 'Table Grid'
    hdr_cells_t = tech_table.rows[0].cells
    headers_t = ["Technical Spec", "Tender Requirement", "L&T Capability / Gap", "Status"]
    for i, name in enumerate(headers_t):
        hdr_cells_t[i].text = name
        set_cell_background(hdr_cells_t[i], "708090") # Slate grey header
        hdr_cells_t[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells_t[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells_t[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_margins(hdr_cells_t[i], top=80, bottom=80, left=100, right=100)

    for check in result.technical_scorecard:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = f"{check.parameter}\n({check.citation})"
        row_cells[1].text = check.tender_requirement
        
        company_text = check.lt_capability
        if check.gap_analysis:
            company_text += f"\nGap: {check.gap_analysis}"
        if check.mitigation_action:
            company_text += f"\nMitigation: {check.mitigation_action}"
        row_cells[2].text = company_text
        
        status_cell = row_cells[3]
        status_cell.text = check.status.value
        
        # Color match status cell
        if check.status.value == "COMPLIANT":
            set_cell_background(status_cell, "D4EDDA")
        elif check.status.value == "RISKY" or check.status.value == "PARTIALLY_COMPLIANT":
            set_cell_background(status_cell, "FFF3CD")
        else:
            set_cell_background(status_cell, "F8D7DA")
            
        for c in row_cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].style.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Footer note
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential - Larsen & Toubro Defence Internal Use Only"
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.style.font.size = Pt(8.5)
    fp.style.font.color.rgb = COLOR_SECONDARY

    doc.save(filepath)

# ==========================================
# 2. Background Processing Worker
# ==========================================
def run_analysis_task(task_id: str, file_path: str, profile_path: str):
    """Background worker method for running 3-stage analysis."""
    analysis_tasks[task_id]["status"] = "Extracting Specs (Stage 1)..."
    try:
        analyzer = TenderAnalyzer()
        
        # Run stage 1-3 full engine
        analysis_tasks[task_id]["status"] = "Running Compliance Engine (Stage 2 & 3)..."
        result = analyzer.analyze_tender(file_path, profile_path)
        
        # Save results locally in memory and disk
        analysis_tasks[task_id]["status"] = "Completed"
        analysis_tasks[task_id]["result"] = result.model_dump()
        
        # Save cache outputs
        with open(OUTPUT_PATH, "w") as f:
            json.dump(result.model_dump(), f, indent=2)
            
    except Exception as e:
        analysis_tasks[task_id]["status"] = "Failed"
        analysis_tasks[task_id]["error"] = str(e)
    finally:
        # Cleanup uploaded local file
        if os.path.exists(file_path):
            os.remove(file_path)

# ==========================================
# 3. Router Endpoints
# ==========================================

@app.post("/analyze")
def analyze_tender_endpoint(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """Receives a PDF/MD file upload and initiates async hybrid analysis."""
    # Generate a unique job/task ID
    task_id = f"job_{int(time.time())}"
    
    # Save the file to temp location
    filename = f"{task_id}_{file.filename}"
    file_path = os.path.join(TEMP_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    analysis_tasks[task_id] = {
        "status": "Initiating",
        "filename": file.filename,
        "result": None,
        "error": None
    }
    
    # Enqueue background task
    background_tasks.add_task(run_analysis_task, task_id, file_path, PROFILE_PATH)
    
    return {"task_id": task_id, "status": "Queued"}

@app.get("/status/{task_id}")
def check_status_endpoint(task_id: str):
    """Queries current status of the analysis task."""
    if task_id not in analysis_tasks:
        raise HTTPException(status_code=404, detail="Task ID not found")
    task_info = analysis_tasks[task_id]
    return {
        "task_id": task_id,
        "status": task_info["status"],
        "error": task_info["error"]
    }

@app.get("/results/{task_id}")
def get_results_endpoint(task_id: str):
    """Retrieves computed scorecard results once task is completed."""
    if task_id not in analysis_tasks:
        raise HTTPException(status_code=404, detail="Task ID not found")
    task_info = analysis_tasks[task_id]
    if task_info["status"] != "Completed":
        raise HTTPException(status_code=400, detail=f"Task is in status: {task_info['status']}")
    return task_info["result"]

@app.get("/profile")
def get_profile():
    """Reads L&T profile JSON data."""
    if not os.path.exists(PROFILE_PATH):
        raise HTTPException(status_code=404, detail="Profile database file not found")
    with open(PROFILE_PATH, "r") as f:
        return json.load(f)

@app.post("/profile")
def update_profile(profile_data: Dict[str, Any]):
    """Saves updated parameters back to company_profile.json."""
    try:
        # Validate schema first using Pydantic model
        validated_profile = CompanyProfile.model_validate(profile_data)
        with open(PROFILE_PATH, "w") as f:
            json.dump(validated_profile.model_dump(), f, indent=2)
        return {"status": "success", "message": "Company profile updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid profile schema: {e}")

@app.get("/export/{task_id}")
def export_docx(task_id: str):
    """Generates and returns the downloadable formatted Word (.docx) report."""
    if task_id not in analysis_tasks:
        raise HTTPException(status_code=404, detail="Task ID not found")
    task_info = analysis_tasks[task_id]
    if task_info["status"] != "Completed":
        raise HTTPException(status_code=400, detail="Cannot export unfinished task")
        
    result_data = BiddingTemplate.model_validate(task_info["result"])
    export_filename = f"L&T_Bidding_Report_{result_data.tender_id}.docx"
    export_path = os.path.join(TEMP_DIR, export_filename)
    
    try:
        build_docx_report(result_data, export_path)
        return FileResponse(
            path=export_path,
            filename=export_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error compiling Word document: {e}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app_backend:app", host="127.0.0.1", port=8000, reload=True)
